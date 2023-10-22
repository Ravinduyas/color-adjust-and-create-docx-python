import os
import zipfile
import subprocess
from PIL import Image, ImageEnhance
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def adjust_image(image_path, output_folder):
    # Open the image
    img = Image.open(image_path)

    # Adjust brightness
    enhancer = ImageEnhance.Brightness(img)
    img = enhancer.enhance(1.5)

    # Adjust contrast
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(1.5)

    # Adjust sharpness
    enhancer = ImageEnhance.Sharpness(img)
    img = enhancer.enhance(5)

    # Generate a random folder name
    folder_name = 'adjusted_images'

    # Create the folder if it doesn't exist
    output_folder = os.path.join(output_folder, folder_name)
    os.makedirs(output_folder, exist_ok=True)

    # Save the adjusted image in the new folder
    file_name = os.path.basename(image_path)
    output_path = os.path.join(output_folder, file_name)
    img.save(output_path)

    return output_path

def process_images(input_folder, output_folder):
    image_files = [f for f in os.listdir(input_folder) if f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif'))]

    adjusted_image_paths = []
    for image_file in image_files:
        image_path = os.path.join(input_folder, image_file)
        adjusted_image_path = adjust_image(image_path, output_folder)
        adjusted_image_paths.append(adjusted_image_path)

    return adjusted_image_paths

def create_word_document(image_paths, output_document):
    doc = Document()

    for image_path in image_paths:
        section = doc.sections[-1]
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.hanging = Pt(720)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Pt(500), height=Pt(700))

    doc.save(output_document)


# Specify the input and output folders
input_folder = 'input_folder' 
output_folder = 'output_folder'  

# Process the images and get their adjusted paths
adjusted_image_paths = process_images(input_folder, output_folder)

# Create a Word document with one image per page and adjust layout
output_document = 'output_document.docx'  
create_word_document(adjusted_image_paths, output_document)

if os.path.exists(output_document):
    # Open the document using the default application
    subprocess.Popen(['start', output_document], shell=True)
else:
    print("The document does not exist at the specified path.")
    
